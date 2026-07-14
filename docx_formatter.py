import io
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def parse_pages(pages_str: str) -> set:
    if pages_str.strip().lower() == "все":
        return "ALL"
    pages = set()
    parts = pages_str.replace(" ", "").split(",")
    for part in parts:
        if "-" in part:
            try:
                start, end = part.split("-")
                if start.isdigit() and end.isdigit():
                    pages.update(range(int(start), int(end) + 1))
            except ValueError:
                pass
        elif part.isdigit():
            pages.add(int(part))
    return pages

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def apply_gost_styles(doc: Document, target_pages="ALL"):
    try:
        if doc.part.numbering_part is not None:
            numbering = doc.part.numbering_part.numbering_definitions._numbering
            for abstract in numbering.findall(qn('w:abstractNum')):
                for lvl in abstract.findall(qn('w:lvl')):
                    lvl_pPr = lvl.find(qn('w:pPr'))
                    if lvl_pPr is None:
                        lvl_pPr = OxmlElement('w:pPr')
                        lvl.append(lvl_pPr)
                    ind = lvl_pPr.find(qn('w:ind'))
                    if ind is None:
                        ind = OxmlElement('w:ind')
                        lvl_pPr.append(ind)
                    for k in list(ind.attrib.keys()):
                        del ind.attrib[k]
                    ind.set(qn('w:left'), '0')
                    ind.set(qn('w:firstLine'), '709')
                    
                    rPr = lvl.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        lvl.append(rPr)
                    
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rFonts.set(qn('w:cs'), 'Times New Roman')
                    
                    sz = rPr.find(qn('w:sz'))
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        rPr.append(sz)
                    sz.set(qn('w:val'), '28')
                    
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is None:
                        szCs = OxmlElement('w:szCs')
                        rPr.append(szCs)
                    szCs.set(qn('w:val'), '28')
                    
                    suff = lvl.find(qn('w:suff'))
                    if suff is None:
                        suff = OxmlElement('w:suff')
                        lvl.append(suff)
                    suff.set(qn('w:val'), 'space')
                    
                    numFmt = lvl.find(qn('w:numFmt'))
                    if numFmt is not None and numFmt.get(qn('w:val')) == 'bullet':
                        lvlText = lvl.find(qn('w:lvlText'))
                        if lvlText is None:
                            lvlText = OxmlElement('w:lvlText')
                            lvl.insert(0, lvlText)
                        lvlText.set(qn('w:val'), '–')
    except Exception as e:
        logging.error(f"Error modifying numbering definitions: {e}")

    if target_pages == "ALL":
        for section in doc.sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.0)
            
            header = section.header
            if not header.paragraphs:
                hp = header.add_paragraph()
            else:
                hp = header.paragraphs[0]
                hp.clear()
            hp.text = "НАЗВАНИЕ КОЛЛЕДЖА, СПЕЦИАЛЬНОСТЬ, ОБОЗНАЧЕНИЕ ДОКУМЕНТА"
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in hp.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.all_caps = True
                
            footer = section.footer
            if not footer.paragraphs:
                fp = footer.add_paragraph()
            else:
                fp = footer.paragraphs[0]
                fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run1 = fp.add_run("НАИМЕНОВАНИЕ ДОКУМЕНТА    ГРУППА    ")
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(14)
            run1.font.all_caps = True
            
            run2 = fp.add_run()
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(14)
            add_page_number(run2)

    page_num = 1
    for block in iter_block_items(doc):
        if target_pages == "ALL" or page_num in target_pages:
            if isinstance(block, Paragraph):
                text = block.text.strip()
                block.paragraph_format.space_before = Pt(0)
                block.paragraph_format.space_after = Pt(0)
                
                has_image = any(run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict') for run in block.runs)
                
                orig_align = block.alignment
                if orig_align is None and block.style and block.style.paragraph_format:
                    orig_align = block.style.paragraph_format.alignment
                    
                is_figure_caption = False
                is_text_heading = False
                is_normal_text = False
                
                is_list = False
                if block.style and 'List' in block.style.name:
                    is_list = True
                elif block._element.pPr is not None and block._element.pPr.numPr is not None:
                    is_list = True
                
                is_heading_style = block.style and 'Heading' in block.style.name
                
                is_fig_by_style = block.style and 'рисунок' in block.style.name.lower()
                if is_fig_by_style and len(text) > 150:
                    is_fig_by_style = False
                is_fig_by_text = text.strip().startswith('Рисунок ') and ' - ' in text
                
                is_table_by_text = text.strip().lower().startswith('таблица')
                
                if has_image:
                    block.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    block.paragraph_format.first_line_indent = Cm(0)
                    block.paragraph_format.line_spacing = 1.0
                    block.paragraph_format.keep_with_next = True
                elif is_fig_by_style or is_fig_by_text or ("Рисунок" in text and any('SEQ' in (n.text or '') for r in block.runs for n in r._element.xpath('.//w:instrText'))):
                    is_figure_caption = True
                    block.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    block.paragraph_format.first_line_indent = Cm(0)
                    block.paragraph_format.space_after = Pt(6)
                    block.paragraph_format.line_spacing = 1.0
                elif is_table_by_text or ("таблица" in text.lower() and any('SEQ' in (n.text or '') for r in block.runs for n in r._element.xpath('.//w:instrText'))):
                    block.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    block.paragraph_format.first_line_indent = Cm(0)
                    block.paragraph_format.left_indent = Cm(0)
                    block.paragraph_format.space_after = Pt(0)
                    block.paragraph_format.line_spacing = 1.0
                elif is_heading_style or orig_align in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
                    if is_heading_style or orig_align == WD_ALIGN_PARAGRAPH.CENTER:
                        is_text_heading = True
                        block.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    block.paragraph_format.first_line_indent = Cm(0)
                    block.paragraph_format.line_spacing = 1.5
                else:
                    is_normal_text = True
                    block.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    block.paragraph_format.line_spacing = 1.5
                    if not is_list:
                        block.paragraph_format.left_indent = Cm(0)
                        block.paragraph_format.first_line_indent = Cm(1.25)
                    else:
                        if block._element.pPr is not None:
                            for old_ind in block._element.pPr.findall(qn('w:ind')):
                                block._element.pPr.remove(old_ind)
                
                for run in block.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    if is_figure_caption:
                        run.font.bold = False
                        run.font.italic = False
                    elif is_text_heading:
                        run.font.bold = True
                        run.font.italic = False
                    elif is_normal_text:
                        run.font.bold = False
                        run.font.italic = False
                    
            elif isinstance(block, Table):
                block.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for row_idx, row in enumerate(block.rows):
                    is_header = (row_idx == 0)
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if is_header:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            para.paragraph_format.line_spacing = 1.0
                            para.paragraph_format.first_line_indent = Cm(0)
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                if is_header:
                                    run.font.bold = True
        
        for elem in block._element.iter():
            if elem.tag.endswith('lastRenderedPageBreak'):
                page_num += 1
            elif elem.tag.endswith('br') and elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                page_num += 1

    if target_pages == "ALL":
        max_w = Cm(17.0)
        max_h = Cm(13.0)
        for shape in doc.inline_shapes:
            if shape.width and shape.height:
                if shape.width > max_w or shape.height > max_h:
                    ratio_w = max_w / shape.width
                    ratio_h = max_h / shape.height
                    ratio = min(ratio_w, ratio_h)
                    if ratio < 1.0:
                        shape.width = int(shape.width * ratio)
                        shape.height = int(shape.height * ratio)
