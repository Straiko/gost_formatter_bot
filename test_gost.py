import sys
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = "/home/root2506/Загрузки/Telegram Desktop/ГОСТ_АЛЛАХ(ТАРАСЮК) (2)_ГОСТ.docx"
doc = Document(doc_path)

errors = []

for i, section in enumerate(doc.sections):
    if round(section.left_margin.cm, 2) != 2.5: errors.append(f"Left margin is {round(section.left_margin.cm, 2)} cm instead of 2.5")
    if round(section.right_margin.cm, 2) != 1.0: errors.append(f"Right margin is {round(section.right_margin.cm, 2)} cm instead of 1.0")
    if round(section.top_margin.cm, 2) != 1.0: errors.append(f"Top margin is {round(section.top_margin.cm, 2)} cm instead of 1.0")
    if round(section.bottom_margin.cm, 2) != 1.0: errors.append(f"Bottom margin is {round(section.bottom_margin.cm, 2)} cm instead of 1.0")

header = doc.sections[0].header
if header.paragraphs:
    hp = header.paragraphs[0]
    if hp.alignment != WD_ALIGN_PARAGRAPH.LEFT: errors.append("Header is not LEFT aligned.")
    for run in hp.runs:
        if run.font.size and run.font.size.pt != 14.0: errors.append(f"Header font size is {run.font.size.pt}")

footer = doc.sections[0].footer
if footer.paragraphs:
    fp = footer.paragraphs[0]
    if fp.alignment != WD_ALIGN_PARAGRAPH.RIGHT: errors.append("Footer is not RIGHT aligned.")
    for run in fp.runs:
        if run.font.size and run.font.size.pt != 14.0: errors.append(f"Footer font size is {run.font.size.pt}")

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text: continue
    
    align = p.alignment
    if align is None and p.style and p.style.paragraph_format:
        align = p.style.paragraph_format.alignment
        
    indent = round(p.paragraph_format.first_line_indent.cm, 2) if p.paragraph_format.first_line_indent else 0.0
    spacing = p.paragraph_format.line_spacing
    
    if "ДЕПАРТАМЕНТ" in text or "МКАГ" in text or "ОТЧЕТ" in text:
        if align != WD_ALIGN_PARAGRAPH.CENTER: errors.append(f"Title '{text[:15]}' not CENTER.")
        if indent != 0.0: errors.append(f"Title '{text[:15]}' has indent.")
    elif "Выполнил обучающийся" in text:
        if align != WD_ALIGN_PARAGRAPH.RIGHT: errors.append(f"Signature '{text[:15]}' not RIGHT.")
    elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        if indent != 1.25: errors.append(f"Para '{text[:15]}' has indent {indent}, should be 1.25.")
        if spacing != 1.5: errors.append(f"Para '{text[:15]}' has spacing {spacing}, should be 1.5.")

for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.size and run.font.size.pt != 12.0:
                        errors.append(f"Table cell '{p.text[:15]}' font size {run.font.size.pt} != 12.")

if not errors:
    print("ALL TESTS PASSED! Document perfectly matches GOST.")
else:
    print("ERRORS FOUND:")
    for e in sorted(set(errors)):
        print("-", e)
