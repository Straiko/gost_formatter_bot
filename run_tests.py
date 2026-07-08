import sys
import io
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bot import apply_gost_styles
try:
    from PIL import Image
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow"])
    from PIL import Image

img = Image.new('RGB', (1600, 1200), color = 'red')
img.save('test_img.png')

doc = Document()
doc.add_paragraph("МИНИСТЕРСТВО").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Выполнил студент").alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph("Нормальный текст абзаца, который должен стать по ширине и с отступом.")
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_picture('test_img.png', width=Cm(20)) # Too big, should be resized
p2 = doc.add_paragraph("Рисунок 1 - Тестовый рисунок")

p3 = doc.add_paragraph("Таблица 1 - Данные")
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = "Заголовок 1"
table.cell(0, 1).text = "Заголовок 2"
table.cell(1, 0).text = "Данные 1"
table.cell(1, 1).text = "Данные 2"

apply_gost_styles(doc, "ALL")

errors = []
sec = doc.sections[0]
if round(sec.left_margin.cm, 2) != 2.5: errors.append(f"Left margin: {sec.left_margin.cm}")
if round(sec.right_margin.cm, 2) != 1.0: errors.append("Right margin")
if round(sec.top_margin.cm, 2) != 1.0: errors.append("Top margin")
if round(sec.bottom_margin.cm, 2) != 1.0: errors.append("Bottom margin")

for p in doc.paragraphs:
    text = p.text.strip()
    if not text: continue
    
    align = p.alignment
    indent = round(p.paragraph_format.first_line_indent.cm, 2) if p.paragraph_format.first_line_indent else 0.0
    spacing = p.paragraph_format.line_spacing
    
    if "МИНИСТЕРСТВО" in text:
        if align != WD_ALIGN_PARAGRAPH.CENTER: errors.append(f"Title alignment error: {align}")
        if indent != 0.0: errors.append(f"Title indent error: {indent}")
    elif "Выполнил студент" in text:
        if align != WD_ALIGN_PARAGRAPH.RIGHT: errors.append(f"Right alignment error: {align}")
        if indent != 0.0: errors.append(f"Right indent error: {indent}")
    elif "Нормальный текст" in text:
        if align != WD_ALIGN_PARAGRAPH.JUSTIFY: errors.append(f"Normal text alignment error: {align}")
        if indent != 1.25: errors.append(f"Normal text indent error: {indent}")
        if spacing != 1.5: errors.append(f"Normal text spacing error: {spacing}")
    elif "Рисунок" in text:
        if align != WD_ALIGN_PARAGRAPH.CENTER: errors.append(f"Figure caption alignment error: {align}")
        if indent != 0.0: errors.append(f"Figure caption indent error: {indent}")
    elif "Таблица" in text:
        if align != WD_ALIGN_PARAGRAPH.LEFT: errors.append(f"Table caption alignment error: {align}")
        if indent != 0.0: errors.append(f"Table caption indent error: {indent}")

for shape in doc.inline_shapes:
    if round(shape.width.cm, 2) > 17.0: errors.append(f"Image width too large: {shape.width.cm}")
    if round(shape.height.cm, 2) > 13.0: errors.append(f"Image height too large: {shape.height.cm}")

if errors:
    print("Errors found:")
    for e in errors: print("-", e)
else:
    print("All tests passed successfully!")
