import sys
from docx import Document

doc_path = "/home/root2506/Загрузки/Telegram Desktop/ГОСТ_АЛЛАХ(ТАРАСЮК) (2).docx"
doc = Document(doc_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs[:30]):
    print(f"Para {i}: text='{p.text[:50]}...', style='{p.style.name}'")

for i, t in enumerate(doc.tables):
    print(f"Table {i}: rows={len(t.rows)}, cols={len(t.columns)}")
