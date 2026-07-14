import os
import io
import logging
from docx import Document

import telebot
from telebot.types import ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove

from docx_formatter import apply_gost_styles, parse_pages
from ai_processor import process_document_with_ai

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

TOKEN = os.getenv("TELEGRAM_TOKEN") or os.getenv("BOT_TOKEN")

if not TOKEN:
    logging.critical("TELEGRAM_TOKEN is missing! Please set it in environment variables.")
    exit(1)

bot = telebot.TeleBot(TOKEN)
user_data = {}

@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, "Привет! Я бот для форматирования документов по ГОСТ. 📄\n\nВы можете:\n1. Отправить мне текст обычным сообщением, и я соберу из него Word-документ.\n2. Отправить файл `.docx`.", reply_markup=ReplyKeyboardRemove())

@bot.message_handler(commands=['cancel'])
def cancel(message):
    if message.chat.id in user_data:
        del user_data[message.chat.id]
    bot.reply_to(message, "Действие отменено.", reply_markup=ReplyKeyboardRemove())

@bot.message_handler(content_types=['text'])
def handle_text(message):
    if message.chat.id in user_data and user_data[message.chat.id].get('state'):
        return
    text = message.text
    doc = Document()
    for paragraph in text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
            
    apply_gost_styles(doc, "ALL")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    bot.send_document(
        chat_id=message.chat.id,
        document=( "ГОСТ_Документ.docx", buffer ),
        caption="Ваш текст успешно структурирован и оформлен по ГОСТу! ✅"
    )

@bot.message_handler(content_types=['document'])
def handle_document(message):
    if not message.document.file_name.endswith('.docx'):
        bot.reply_to(message, "Пожалуйста, отправьте файл в формате .docx (Word).")
        return
        
    file_info = bot.get_file(message.document.file_id)
    downloaded_file = bot.download_file(file_info.file_path)
    
    user_data[message.chat.id] = {
        'file_bytes': downloaded_file,
        'file_name': message.document.file_name,
        'state': 'WAITING_FOR_PAGES'
    }
    
    markup = ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    markup.add(KeyboardButton("Все"))
    bot.reply_to(
        message,
        "К каким страницам применить форматирование ГОСТ (шрифты, абзацы)?\n\n"
        "Напишите номера страниц через запятую или дефис (например: `1, 3, 5-10`).\n"
        "Либо нажмите «Все», чтобы отформатировать весь документ целиком.\n\n"
        "*(Внимание: Номера страниц рассчитываются самим редактором Word. Частичное форматирование может сбиться, если текст съедет на другие страницы. Поля документа и колонтитулы останутся нетронутыми!)*.",
        parse_mode='Markdown',
        reply_markup=markup
    )
    bot.register_next_step_handler(message, process_document_pages)

def process_document_pages(message):
    if message.text == '/cancel':
        return cancel(message)
    pages_str = message.text
    target_pages = parse_pages(pages_str)
    
    if not target_pages:
        bot.reply_to(message, "Не удалось распознать номера страниц. Попробуйте еще раз (например: 1, 3-5) или напишите 'Все'.")
        bot.register_next_step_handler(message, process_document_pages)
        return
        
    user_data[message.chat.id]['target_pages'] = target_pages
    user_data[message.chat.id]['state'] = 'WAITING_FOR_AI'
    
    markup = ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    markup.add(KeyboardButton("0"), KeyboardButton("1"), KeyboardButton("2"), KeyboardButton("3"))
    bot.reply_to(
        message,
        "Хотите ли вы, чтобы ИИ проанализировал текст и автоматически сгенерировал списки?\n"
        "Укажите максимальное количество списков на главу (например: 0 — не нужно, 1, 2...)",
        reply_markup=markup
    )
    bot.register_next_step_handler(message, process_ai_lists)

def process_ai_lists(message):
    if message.text == '/cancel':
        return cancel(message)
    try:
        lists_per_chapter = int(message.text)
    except ValueError:
        bot.reply_to(message, "Пожалуйста, введите число (0, 1, 2...)")
        bot.register_next_step_handler(message, process_ai_lists)
        return

    status_message = bot.reply_to(message, "⏳ Обрабатываю документ (это может занять около минуты)...", reply_markup=ReplyKeyboardRemove())
    
    try:
        file_bytes = user_data[message.chat.id]['file_bytes']
        doc = Document(io.BytesIO(file_bytes))
        target_pages = user_data[message.chat.id].get('target_pages', 'ALL')
        
        if lists_per_chapter > 0:
            doc = process_document_with_ai(doc, lists_per_chapter)
        
        apply_gost_styles(doc, target_pages)
        
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        caption_text = "Ваш документ отформатирован! ✅"
        if target_pages != "ALL":
            caption_text += "\nШрифты и абзацы изменены только на выбранных страницах."
            
        bot.send_document(
            chat_id=message.chat.id,
            document=(f"ГОСТ_{user_data[message.chat.id]['file_name']}", output_buffer),
            caption=caption_text
        )
        
        bot.delete_message(chat_id=message.chat.id, message_id=status_message.message_id)
        
    except Exception as e:
        logging.error(f"Error processing file for {message.chat.id}: {e}", exc_info=True)
        try:
            bot.reply_to(message, "❌ Произошла ошибка при обработке файла.")
        except:
            pass
        
    if message.chat.id in user_data:
        del user_data[message.chat.id]

if __name__ == "__main__":
    logging.info("Starting GOST Formatter Bot...")
    bot.infinity_polling(timeout=10, long_polling_timeout=5)
