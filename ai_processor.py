import os
import json
import re
import logging
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI

def process_document_with_ai(doc: Document, lists_per_chapter: int):
    if lists_per_chapter <= 0:
        return doc
        
    client = OpenAI(
        api_key=os.getenv("PROXYAPI_KEY"),
        base_url="https://api.proxyapi.ru/openai/v1"
    )
    
    text_blocks = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 15]
    chunk_size = 60
    
    for i in range(0, len(text_blocks), chunk_size):
        chunk = "\\n\\n".join(text_blocks[i:i+chunk_size])
        prompt = f"""You are an expert editor formatting Russian text according to GOST. 
Here is a text chunk. Your task is to find places where formatting should be fixed. 
This includes:
1) Lists: Sequences of paragraphs that are clearly list items but missing bullet points. Add dashes ("- ") or numbers ("1) "). 
2) Headings: You MUST identify ALL headings and wrap them in <HEADING> and </HEADING> tags.

CRITICAL RULES:
- "original_paragraphs" MUST contain EVERY SINGLE PARAGRAPH from the original text that you are replacing.
- "new_paragraphs" MUST contain only the transformed versions of the exact text in "original_paragraphs". DO NOT hallucinate.
- The introductory sentence before a list MUST end with a COLON (:).
- For dashed lists ("- "): The text of every item MUST start with a LOWERCASE letter. Every item MUST end with a semicolon (";"), EXCEPT the very last item which MUST end with a period (".").

Return ONLY a JSON object with a single key "replacements", which is an array of objects.
"original_paragraphs": ["Exact text 1", "Exact text 2"],
"new_paragraphs": ["Transformed text 1", "Transformed text 2"]

If no paragraphs need conversion, return {{"replacements": []}}.
Text chunk:
{chunk}"""
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"}
            )
            data = json.loads(response.choices[0].message.content)
            replacements = data.get("replacements", [])
            
            _apply_replacements(doc, replacements)
        except Exception as e:
            logging.error(f"ProxyAPI error: {e}")
            
    return doc

def _apply_replacements(doc, replacements):
    for rep in replacements:
        orig_paras = rep.get("original_paragraphs", [])
        if "original_paragraph" in rep and not orig_paras:
            orig_paras = [rep["original_paragraph"]]
            
        new_paras = rep.get("new_paragraphs", [])
        if not orig_paras or not new_paras: continue
        
        orig_paras = [p.strip() for p in orig_paras if p.strip()]

        def normalize(t):
            return re.sub(r'\W+', '', t.lower())

        all_paras = doc.paragraphs
        for idx in range(len(all_paras) - len(orig_paras) + 1):
            match = True
            for j in range(len(orig_paras)):
                actual_norm = normalize(all_paras[idx+j].text)
                orig_norm   = normalize(orig_paras[j])
                key = orig_norm[:60]
                if not key:
                    match = False; break
                if actual_norm != orig_norm and not actual_norm.startswith(key):
                    match = False
                    break
            
            if match:
                elements_to_delete = []
                for j in range(1, len(orig_paras)):
                    elements_to_delete.append(all_paras[idx+j]._element)
                
                for p_idx, item in enumerate(new_paras):
                    item = item.strip()
                    is_heading = False
                    
                    if item.startswith("<HEADING>") and item.endswith("</HEADING>"):
                        item = item[9:-10].strip()
                        is_heading = True
                        
                    if p_idx == 0:
                        all_paras[idx].text = item
                        current_p = all_paras[idx]
                    else:
                        list_type = None
                        if item.startswith("-"):
                            item = item.lstrip("- ").strip()
                            list_type = 'bullet'
                        elif re.match(r'^\d+[\)\.]', item):
                            item = re.sub(r'^\d+[\)\.]\s*', '', item).strip()
                            list_type = 'number'

                        if list_type:
                            target_styles = ['List Bullet' if list_type == 'bullet' else 'List Number']
                            new_p = None
                            for s_name in target_styles:
                                try:
                                    new_p = doc.add_paragraph(item, style=s_name)
                                    break
                                except KeyError:
                                    pass
                            if not new_p:
                                for s in doc.styles:
                                    if s.type == 1:
                                        name = s.name.lower()
                                        if list_type == 'bullet' and ('bullet' in name or 'маркиров' in name):
                                            new_p = doc.add_paragraph(item, style=s.name)
                                            break
                                        if list_type == 'number' and ('number' in name or 'нумеров' in name):
                                            new_p = doc.add_paragraph(item, style=s.name)
                                            break
                            if not new_p:
                                try:
                                    new_p = doc.add_paragraph(item, style='List Paragraph')
                                except KeyError:
                                    new_p = doc.add_paragraph(item)
                                pPr = new_p._p.get_or_add_pPr()
                                numPr = OxmlElement('w:numPr')
                                ilvl = OxmlElement('w:ilvl')
                                ilvl.set(qn('w:val'), '0')
                                numId = OxmlElement('w:numId')
                                numId.set(qn('w:val'), '1' if list_type == 'bullet' else '2')
                                numPr.append(ilvl)
                                numPr.append(numId)
                                pPr.append(numPr)

                            # Indentation is now handled globally by docx_formatter.py
                        else:
                            new_p = doc.add_paragraph(item)
                            
                        current_p._element.addnext(new_p._element)
                        current_p = new_p

                    if is_heading:
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        current_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in current_p.runs:
                            r.font.bold = True
                
                for elem in elements_to_delete:
                    parent = elem.getparent()
                    if parent is not None:
                        parent.remove(elem)
                    
                break
